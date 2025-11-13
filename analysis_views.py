import pdfplumber 
import re 
import nltk
from nltk.tokenize import word_tokenize 
from nltk.corpus import stopwords 
from nltk.stem import WordNetLemmatizer 

from sklearn.feature_extraction.text import TfidfVectorizer 
from sklearn.metrics.pairwise import cosine_similarity 

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

from job_analysis.models import Resume, ResumeAnalysis, JobDescription, CustomUser
from job_analysis.utils import auth_user, jwt_decode

nltk.download('punkt')
nltk.download('punkt_tab')
nltk.download('wordnet')
nltk.download('stopwords')

def preprocess_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z\s]', '', text)
    words = word_tokenize(text, language='english')
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]
    lemmatizer = WordNetLemmatizer()
    words = [lemmatizer.lemmatize(word) for word in words]
    return ' '.join(words)

def extract_text_from_file(file):
    """
    Extract text from either PDF or DOCX file.
    
    Args:
        file: File object or file path
        
    Returns:
        str: Extracted and preprocessed text
    """
    text = ""
    
    try:
        # Handle InMemoryUploadedFile from Django
        if hasattr(file, 'temporary_file_path'):
            file_path = file.temporary_file_path()
            file_extension = file.name.split('.')[-1].lower()
        else:
            file_path = file
            file_extension = str(file).split('.')[-1].lower()
            
        if file_extension == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or "" + "\n"
        
        elif file_extension in ['docx', 'doc']:
            try:
                from docx import Document
                import io
                
                if hasattr(file, 'read'):
                    # Handle file-like object
                    doc = Document(io.BytesIO(file.read()))
                else:
                    # Handle file path
                    doc = Document(file_path)
                    
                for para in doc.paragraphs:
                    text += para.text + "\n"
                    
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                        
            except ImportError:
                raise ImportError("python-docx package is required to process DOCX files. Please install it using 'pip install python-docx'")
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Please provide a PDF or DOCX file.")

        if not text.strip():
            raise ValueError(f"No text extracted from the file. Please check the file.")

        return preprocess_text(text)

    except Exception as e:
        print(f"Error extracting text from file: {e}")
        return None

def extract_skills(text):
    predefined_skills = [
        # Technical Skills (General)
        "Python", "Java", "C++", "JavaScript", "SQL", "Machine Learning", 
        "Data Science", "Django", "Flask", "HTML", "CSS", "React", "Node.js", 
        "AWS", "Azure", "Docker", "Kubernetes", "Git", "PostgreSQL", "MongoDB", 
        "Agile", "Leadership", "Teamwork",
        
        # AI/ML & Data Science
        "TensorFlow", "PyTorch", "OpenCV", "NLP", "Computer Vision",
        "Deep Learning", "Neural Networks", "Reinforcement Learning",
        "Natural Language Processing", "Text Analytics",
        "Data Mining", "Predictive Analytics",
        "Big Data Analytics", "Data Warehousing",
        "ETL", "Data Pipelines", "Data Integration",
        "Machine Learning Operations", "MLOps",
        
        # Cloud & Infrastructure
        "AWS Lambda", "AWS S3", "AWS EC2", "AWS RDS",
        "Azure Functions", "Azure Blob Storage", "Azure VMs",
        "Google Cloud Platform", "GCP Cloud Storage",
        "Cloud Architecture", "Cloud Security",
        "Infrastructure as Code", "Infrastructure Automation",
        "Terraform", "Ansible", "SaltStack",
        
        # DevOps & CI/CD
        "CI/CD", "Continuous Integration", "Continuous Deployment",
        "Jenkins", "GitLab CI", "GitHub Actions",
        "Docker Compose", "Docker Swarm", "Kubernetes",
        "Kafka", "RabbitMQ", "Message Queues",
        "Redis", "Memcached", "Caching",
        "Monitoring", "Logging", "Debugging",
        "Performance Optimization", "Load Testing",
        "Chaos Engineering", "Site Reliability Engineering",
        
        # Web Development
        "REST API", "GraphQL", "Microservices", "Serverless",
        "Next.js", "Nuxt.js", "Gatsby",
        "Vue.js", "Angular", "Svelte",
        "TypeScript", "GraphQL", "Apollo",
        "Progressive Web Apps", "Web Components",
        "Web Accessibility", "Performance Optimization",
        
        # Mobile Development
        "React Native", "Flutter", "Swift", "Kotlin",
        "Android", "iOS", "Xamarin", "Ionic",
        "Mobile UI/UX", "Mobile Performance",
        "Push Notifications", "Location Services",
        
        # Game Development
        "Unity", "Unreal Engine", "Godot",
        "Game Physics", "Game AI", "Game Graphics",
        "Game Design", "Level Design", "Game Testing",
        
        # AR/VR & 3D
        "ARKit", "ARCore", "WebXR",
        "Three.js", "WebGL", "WebGPU",
        "3D Modeling", "3D Animation", "3D Rendering",
        "Virtual Reality", "Augmented Reality",
        
        # Security
        "Cyber Security", "Penetration Testing", "Ethical Hacking",
        "Network Security", "Application Security",
        "Identity and Access Management", "IAM",
        "Security Compliance", "Security Auditing",
        "Security Architecture",
        
        # Dev Tools & IDEs
        "Visual Studio Code", "IntelliJ IDEA", "PyCharm",
        "Eclipse", "NetBeans", "Sublime Text",
        "Postman", "JMeter", "Selenium",
        "Jira", "Trello", "Asana",
        
        # UI/UX Design
        "UI/UX Design", "Adobe Creative Suite", "Figma", "Sketch",
        "Adobe XD", "InVision", "Zeplin",
        "Wireframing", "Prototyping", "User Testing",
        "Responsive Design", "Accessibility Design",
        
        # Data Visualization
        "Tableau", "Power BI", "Matplotlib", "Seaborn",
        "Plotly", "D3.js", "Bokeh",
        "Data Storytelling", "Dashboard Design",
        
        # Version Control
        "Git", "GitLab", "Bitbucket", "GitHub Actions",
        "Branch Management", "Code Review",
        "Version Control Best Practices",
        
        # Testing
        "JUnit", "Selenium", "Postman", "Load Testing",
        "Test-Driven Development", "Behavior-Driven Development",
        "Unit Testing", "Integration Testing",
        "Performance Testing", "Security Testing",
        
        # Database
        "Database Design", "Database Optimization",
        "SQL", "NoSQL", "MongoDB", "Cassandra",
        "Redis", "Elasticsearch", "Neo4j",
        
        # Network & Systems
        "Network Administration", "System Administration",
        "Linux Administration", "Windows Administration",
        "Network Security", "Network Troubleshooting",
        
        # IoT & Embedded Systems
        "IoT", "Embedded Systems", "Arduino", "Raspberry Pi",
        "Microcontrollers", "Real-time Systems",
        "Embedded Software", "Firmware Development",
        
        # Engineering Skills
        "Thermodynamics", "Fluid Mechanics", "Structural Analysis",
        "Circuit Design", "Control Systems", "Signal Processing",
        "Material Science", "Process Engineering", "Chemical Engineering",
        "Aerodynamics", "Structural Design", "Mechanical Design",
        "Electrical Systems", "Power Systems", "Electronics",
        "Civil Engineering", "Construction Management",
        "Chemical Process Design", "Process Optimization",
        "Biomedical Instrumentation", "Medical Devices",
        "Environmental Systems", "Waste Management",
        "Materials Testing", "Manufacturing Processes",
        "Industrial Engineering", "Operations Research",
        
        # Medical Skills
        "Clinical Research", "Patient Care", "Medical Diagnosis", "Treatment Planning",
        "Medical Records Management", "Healthcare Documentation", "Patient Assessment",
        "Emergency Medicine", "Surgery", "Anesthesia", "Radiology", "Pathology",
        "Pharmacology", "Medical Ethics", "Infection Control", "Sterilization",
        "Medical Equipment Operation", "Diagnostic Testing",
        "Nursing Care", "Physiotherapy", "Pharmacy Management",
        "Optometry", "Medical Imaging", "Medical Laboratory",
        
        # Business Skills
        "Financial Analysis", "Budgeting", "Cost Management",
        "Marketing Strategy", "Market Research", "Sales Management",
        "Human Resource Management", "Recruitment", "Training",
        "Economic Analysis", "Financial Planning",
        "Accounting", "Taxation", "Auditing",
        "Business Strategy", "Entrepreneurship",
        "Customer Relationship Management",
        "Supply Chain Management", "Logistics",
        
        # Science Skills
        "Physics", "Chemistry", "Biology", "Mathematics", "Statistics",
        "Research Methodology", "Experimental Design",
        "Data Analysis", "Scientific Computing",
        "Environmental Science", "Ecology",
        "Biotechnology", "Genetic Engineering",
        "Biochemistry", "Molecular Biology",
        "Computational Science", "Scientific Programming",
        
        # Arts Skills
        "Literary Analysis", "Historical Research",
        "Political Science", "Public Policy",
        "Sociology", "Anthropology",
        "Psychological Research", "Behavioral Science",
        "Journalism", "Media Studies",
        "Art History", "Visual Arts",
        "Performance Arts", "Theatre Production",
        
        # Law Skills
        "Legal Research", "Legal Writing",
        "Contract Law", "Corporate Law",
        "Intellectual Property", "Patent Law",
        "Criminal Law", "Civil Law",
        "Legal Advocacy", "Legal Ethics",
        
        # Education Skills
        "Teaching Methodology", "Curriculum Design",
        "Educational Technology", "Learning Assessment",
        "Pedagogy", "Educational Research",
        "Classroom Management", "Student Development",
        
        # Professional Skills
        "Accounting", "Auditing", "Taxation",
        "Legal Practice", "Legal Consultation",
        "Architecture", "Urban Planning",
        "Surveying", "Construction Management",
        
        # Specialized Skills
        "Data Science", "Machine Learning", "AI Development",
        "Digital Marketing", "SEO", "Social Media",
        "Hospitality Management", "Food Service",
        "Event Planning", "Tourism Management",
        "Social Work", "Community Development",
        "Public Administration", "Policy Analysis",
        "Sports Science", "Physical Training",
        "Music Theory", "Composition",
        "Dance Choreography", "Theatre Direction",
        
        # Non-Technical Skills
        "Communication", "Problem Solving", "Time Management", "Adaptability",
        "Critical Thinking", "Decision Making", "Conflict Resolution",
        "Project Management", "Budget Management", "Risk Management",
        "Customer Service", "Negotiation", "Presentation", "Research",
        "Data Analysis", "Report Writing", "Documentation", "Quality Control",
        "Team Leadership", "Mentoring", "Stress Management", "Multi-tasking",
        "Attention to Detail", "Creativity", "Strategic Thinking",
        "Business Development", "Marketing", "Sales",
        
        # Sports Skills
        "Football", "Cricket", "Basketball", "Volleyball", "Badminton",
        "Tennis", "Swimming", "Athletics", "Track and Field",
        "Gymnastics", "Martial Arts", "Karate", "Taekwondo",
        "Yoga", "Aerobics", "Fitness Training",
        "Team Sports", "Individual Sports", "Sports Leadership",
        "Sportsmanship", "Sports Strategy",
        
        # Dance Skills
        "Bharatanatyam", "Kathak", "Kuchipudi", "Odissi",
        "Ballet", "Hip-Hop", "Contemporary", "Jazz",
        "Salsa", "Bhangra", "Garba", "Folk Dance",
        "Choreography", "Performance", "Stage Presence",
        "Dance Technique", "Rhythm", "Expression",
        
        # Singing Skills
        "Classical Singing", "Carnatic Music", "Hindustani Music",
        "Western Classical", "Pop Singing", "Jazz Singing",
        "Vocal Training", "Sight Reading", "Music Theory",
        "Instrumental", "Guitar", "Piano", "Keyboard",
        "Voice Modulation", "Breathing Techniques",
        "Stage Performance", "Audition Preparation",
        "Music Production", "Recording", "Composition"
    ]
    text = text.lower()
    matching_skills = [skill for skill in predefined_skills if skill.lower() in text]
    return matching_skills

def extract_education(text):
    predefined_degrees = [
        # Bachelor's Degrees
        "bachelor", "bachelor's", "bachelor of", "b.sc", "b.com", "bba", "bca",
        "bba", "bsc", "btech", "be", "bpharm", "bhm", "bba", "bsw",
        "bba", "bcom", "bba", "bba", "bba", "bba",
        
        # Master's Degrees
        "master", "master's", "master of", "msc", "mcom", "mba", "mca",
        "mtech", "me", "mpharm", "mhm", "msw", "mcom", "mca",
        "mca", "mca", "mca", "mca",
        
        # Doctoral Degrees
        "phd", "doctorate", "doctor of philosophy",
        
        # Associate Degrees
        "associate", "associate's", "associate degree",
        
        # Engineering Degrees
        "computer science", "information technology", "mechanical engineering",
        "electrical engineering", "civil engineering", "chemical engineering",
        "aerospace engineering", "biomedical engineering", "environmental engineering",
        "materials science", "software engineering", "industrial engineering",
        
        # Medical Degrees
        "mbbs", "md", "ms", "dm", "mch", "dental", "dentistry", "pharmacy",
        "nursing", "physiotherapy", "optometry", "medicine",
        
        # Business Degrees
        "business administration", "management", "commerce", "economics",
        "finance", "accountancy", "marketing", "human resources",
        
        # Science Degrees
        "physics", "chemistry", "biology", "mathematics", "statistics",
        "environmental science", "biotechnology", "biochemistry",
        "computer science", "information technology",
        
        # Arts Degrees
        "english", "history", "political science", "sociology",
        "psychology", "philosophy", "journalism", "mass communication",
        "fine arts", "visual arts", "performing arts",
        
        # Law Degrees
        "llb", "llm", "law", "juris doctor",
        
        # Education Degrees
        "education", "teaching", "pedagogy",
        
        # Professional Degrees
        "chartered accountant", "ca", "cost accountant", "cs",
        "company secretary", "management accountant", "lawyer",
        "architect", "urban planner", "surveyor",
        
        # Specialized Degrees
        "data science", "artificial intelligence", "machine learning",
        "data analytics", "digital marketing", "hospitality management",
        "hospital management", "hotel management", "social work",
        "public administration", "public health", "environmental management",
        "sports science", "physical education", "music",
        "dance", "theatre", "film studies"
    ]
    text = text.lower()
    matching_degrees = [degree for degree in predefined_degrees if degree in text]
    return matching_degrees

def extract_experience(text):
    experience_years = re.findall(r'(\d+)\s*(?:years?|yrs?)\s*(?:of)?\s*experience', text, re.IGNORECASE)
    if experience_years:
        return max(map(int, experience_years))  # Take the maximum years of experience found
    return 0  # Default to 0 if no experience is found

def calculate_similarity(resume_text, jd_text):
    try:
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
        similarity_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        match_percentage = similarity_matrix[0][0] * 100

        resume_skills = extract_skills(resume_text)
        jd_skills = extract_skills(jd_text)

        matching_skills = list(set(resume_skills) & set(jd_skills))
        missing_skills = list(set(jd_skills) - set(resume_skills))
        extra_skills = list(set(resume_skills) - set(jd_skills))
        skills_match_percentage = len(matching_skills) / len(jd_skills) * 100 if jd_skills else 0

        # Education Match Percentage
        resume_education = extract_education(resume_text)
        jd_education = extract_education(jd_text)
        education_match_percentage = len(set(resume_education) & set(jd_education)) / len(jd_education) * 100 if jd_education else 0

        # Experience Match Percentage
        resume_experience = extract_experience(resume_text)
        jd_experience = extract_experience(jd_text)
        experience_match_percentage = min(resume_experience / jd_experience, 1) * 100 if jd_experience else 0

        analysis_results = {
            "overall_match_percentage": match_percentage,
            "skills_match_percentage": skills_match_percentage,
            "education_match_percentage": education_match_percentage,
            "experience_match_percentage": experience_match_percentage,
            "matching_skills": matching_skills,
            "missing_skills": missing_skills,
            "extra_skills": extra_skills,
            "job_description_summary": jd_text[:300],
            "resume_summary": resume_text[:300],
            "analysis_details": "Further analysis can be done on experience and education matching."
        }

        return analysis_results
    except Exception as e:
        print(f"Error calculating similarity: {e}")
        return None

@csrf_exempt
@require_http_methods(["POST"])
def analyze_resume(request):
    bearer = request.headers.get('Authorization')
    if not bearer:
        return JsonResponse({'success': False, 'message': 'Authentication header is required.'}, status=401)
    
    token = bearer.split()[1]
    if not auth_user(token):
        return JsonResponse({'success': False, 'message': 'Invalid token data.'}, status=401)
    
    decoded_token = jwt_decode(token)
    user_email = decoded_token.get('email')

    try:
        user = CustomUser.objects.get(email=user_email)
    except CustomUser.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'User not found.'}, status=404)
    
    resume_file = request.FILES.get("resume_pdf")
    jd_file = request.FILES.get("job_description_pdf")
    jd_text = request.POST.get("job_description_text")

    if not resume_file:
        return JsonResponse({"success": False, "message": "Resume PDF is required"}, status=400)

    if not (jd_file or jd_text):
        return JsonResponse({"success": False, "message": "Job Description PDF or Text is required"}, status=400)

    resume_text = extract_text_from_file(resume_file)
    if resume_text is None:
        return JsonResponse({"success": False, "message": "Error extracting text from Resume file"}, status=400)

    resume_skills = extract_skills(resume_text)
    resume_education = extract_education(resume_text)
    resume_experience = extract_experience(resume_text)

    try:
        resume = Resume.objects.get(user=user)
    except Resume.DoesNotExist:
        resume = Resume(user=user)
    
    if resume_file:
        resume.resume_file = resume_file
    resume.summary = resume_text[:300]
    resume.skills = ', '.join(resume_skills)
    resume.education = ', '.join(resume_education)
    resume.experience = str(resume_experience)
    resume.save()

    if jd_file:
        jd_text = extract_text_from_file(jd_file)
        if jd_text is None:
            return JsonResponse({"success": False, "message": "Error extracting text from Job Description file"}, status=400)

    jd_skills = extract_skills(jd_text)
    jd_experience = extract_experience(jd_text)
    jd_education = extract_education(jd_text)

    title = jd_text.strip()[:255]
    if not title:
        title = "Job Description"

    jd_instance = JobDescription.objects.create(
        user=user,
        title=title,
        description=jd_text,
        skills_required=', '.join(jd_skills),
        experience_required=str(jd_experience),
    )

    analysis_results = calculate_similarity(resume_text, jd_text)
    if analysis_results is None:
        return JsonResponse({"success": False, "message": "Internal Server Error"}, status=500)

    ResumeAnalysis.objects.create(
        user=user,
        job_description=jd_instance,
        resume=resume,
        match_percentage=analysis_results['overall_match_percentage'],
        missing_skills=', '.join(analysis_results['missing_skills']),
        extra_skills=', '.join(analysis_results['extra_skills']),
        analysis_details=analysis_results['analysis_details'],
    )

    return JsonResponse({"success": True, "message": "Analysis completed successfully", **analysis_results}, status=200)